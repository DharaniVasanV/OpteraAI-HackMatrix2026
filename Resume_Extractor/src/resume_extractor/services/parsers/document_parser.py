import io
import os
from typing import Tuple
from loguru import logger
import docx
import pypdf
from PIL import Image
import pytesseract

from src.resume_extractor.core.config import settings

# Configure tesseract path if binary exists on Windows
tesseract_path = getattr(
    settings, "TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)
if os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path


class DocumentParserService:
    """
    Parser service capable of extracting text from PDF, DOCX, TXT, and image files.
    """

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extracts text and hyperlink URLs from PDF bytes using PyMuPDF (fitz)."""
        import fitz
        import re
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_pages = []
            for page in doc:
                page_text = page.get_text() or ""
                extracted_urls = []

                # 1. Regex search for explicit URLs in extracted text
                urls_in_text = re.findall(r'https?://[^\s<>"\'\)]+', page_text)
                extracted_urls.extend([u.strip() for u in urls_in_text if u.strip()])

                # 2. Extract embedded links using PyMuPDF
                links = page.get_links()
                for link in links:
                    if link.get("kind") == fitz.LINK_URI:
                        uri = link.get("uri")
                        if uri and isinstance(uri, str):
                            extracted_urls.append(uri.strip())

                if extracted_urls:
                    unique_urls = list(dict.fromkeys(extracted_urls))
                    page_text += "\n\nEmbedded Links & URLs:\n" + "\n".join(unique_urls)

                if page_text.strip():
                    text_pages.append(page_text)
            
            return "\n".join(text_pages)
        except Exception as exc:
            logger.warning(f"fitz extraction failed: {exc}")
            return ""



    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extracts text and embedded hyperlink URLs from DOCX bytes using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)

            # Extract embedded hyperlink targets from DOCX document relationships
            docx_urls = []
            if hasattr(doc.part, "rels"):
                for rel in doc.part.rels.values():
                    target = getattr(rel, "target_ref", "")
                    if target and (target.startswith("http://") or target.startswith("https://")):
                        docx_urls.append(target.strip())

            if docx_urls:
                unique_urls = list(dict.fromkeys(docx_urls))
                full_text.append("\nEmbedded Links & URLs:\n" + "\n".join(unique_urls))

            return "\n".join(full_text)
        except Exception as exc:
            logger.warning(f"python-docx extraction failed: {exc}")
            return ""


    @staticmethod
    def extract_text_from_image(file_bytes: bytes) -> str:
        """Extracts text from image bytes using pytesseract OCR."""
        try:
            image = Image.open(io.BytesIO(file_bytes))
            return pytesseract.image_to_string(image)
        except Exception as exc:
            logger.warning(f"pytesseract OCR extraction failed: {exc}")
            return ""

    @classmethod
    def parse_document_with_method(cls, filename: str, file_bytes: bytes) -> Tuple[str, str]:
        """
        Extracts plain text and returns a tuple: (extracted_text, method_used).
        Methods: 'pypdf', 'python-docx', 'tesseract_ocr', 'txt'.
        """
        ext = filename.lower().split(".")[-1] if "." in filename else ""
        extracted_text = ""
        method_used = "unknown"

        if ext == "pdf":
            extracted_text = cls.extract_text_from_pdf(file_bytes)
            method_used = "pypdf"
        elif ext in ["docx", "doc"]:
            extracted_text = cls.extract_text_from_docx(file_bytes)
            method_used = "python-docx"
        elif ext in ["png", "jpg", "jpeg", "tiff", "bmp"]:
            extracted_text = cls.extract_text_from_image(file_bytes)
            method_used = "tesseract_ocr"
        elif ext == "txt":
            method_used = "txt"
            try:
                extracted_text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                extracted_text = file_bytes.decode("latin-1", errors="ignore")
        else:
            raise ValueError(f"Unsupported file format: .{ext}")

        # Fallback to OCR if PDF contained no selectable text (scanned PDF)
        if ext == "pdf" and not extracted_text.strip():
            logger.info("PDF contained no selectable text. Attempting OCR fallback using Pytesseract...")
            extracted_text = cls.extract_text_from_image(file_bytes)
            method_used = "tesseract_ocr_fallback"

        return extracted_text.strip(), method_used

    @classmethod
    def parse_document(cls, filename: str, file_bytes: bytes) -> str:
        """
        Main entry point returning extracted text string.
        """
        text, _ = cls.parse_document_with_method(filename, file_bytes)
        return text
